from typing import List

from llm.llm import parse_pdf_to_json
from database.db import save_tender_to_db
import os
import re
from processor.converter import convert_files_to_pdf
from processor.file_utils import delete_files, rename_all_files_in_folder
import shutil
import hashlib


# grabs keywords from a text file.file format: each keyword on a new line.empty lines and lines starting with # are ignored.
def load_keywords_from_file(filename: str = "file_filter_keywords.txt") -> List[str]:
    keywords = []
    
    # checking the existence of the file
    if not os.path.exists(filename):
        print(f"File {filename} not found. Using default keywords.")
        return get_default_keywords()
    
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            for line_num, line in enumerate(f, 1):
                line = line.strip()
                
                # skip empty lines and comments
                if not line or line.startswith('#'):
                    continue
                
                keywords.append(line)
        
        if not keywords:
            print(f"File {filename} contains no keywords. Using defaults.")
            return get_default_keywords()
        
        print(f"Loaded {len(keywords)} keywords from {filename}")
        return keywords
        
    except Exception as e:
        print(f"Error reading {filename}: {e}")
        print("Using default keywords.")
        return get_default_keywords()

# spits out a standard list of keywords (if the file is not found).
def get_default_keywords() -> List[str]:
    return [
        "описание объекта закупки",
        "ООЗ",
        "наименование объекта закупки",
        "технические характеристики",
        "ТЗ",
        "техническое задание",
        "требования к товару",
        "характеристики товара",
        "функциональные характеристики",
        "показатели товара",
        "спецификация",
        "перечень поставляемого товара",
        "соответствие товара",
        "КТРУ",
        "каталог товаров работ услуг"
    ]

keywords: List[str] = load_keywords_from_file()

# regular expression for quick search (case insensitive)
KEYWORDS_PATTERN = re.compile(
    r"(" + "|".join(re.escape(kw) for kw in keywords) + r")",
    re.IGNORECASE
)

# peeks at if the file contains at least one keyword.supported formats: .txt, .docx, .pdf, .xlsx, .odt, .ods, .rtf
def check_keywords_in_file(file_path: str, keywords_pattern: re.Pattern = KEYWORDS_PATTERN) -> bool:
    try:
        ext = os.path.splitext(file_path)[1].lower()
        
        # text files and html
        if ext in ['.txt', '.html', '.htm', '.xml', '.csv', '.md', '.rst']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(1024 * 1024)  # read the first 1mb for speed
            return bool(keywords_pattern.search(content))
        
        # docx
        elif ext == '.docx':
            try:
                from docx import Document
                doc = Document(file_path)
                text = ' '.join([para.text for para in doc.paragraphs])
                # also check the tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += ' ' + cell.text
                match = keywords_pattern.search(text[:1024*1024])
                if match:
                    print(f"keyword: {match.group(0)}")
                    return True
                else:
                    return False
            except ImportError:
                print("python-docx not installed. Run: pip install python-docx")
                return False
        
        # pdf
        elif ext == '.pdf':
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                text = ''
                for page in reader.pages[:10]:  # first 10 pages
                    text += page.extract_text() or ''
                return bool(keywords_pattern.search(text[:1024*1024]))
            except ImportError:
                print("pypdf not installed. Run: pip install pypdf")
                return False
        
        # excel
        elif ext in ['.xlsx', '.xls']:
            try:

                
                # loading the workbook (.xlsx only, .xls is not supported by openpyxl)
                wb = load_workbook(file_path, read_only=True, data_only=True)
                
                text_parts = []
                char_count = 0
                max_chars = 1024 * 1024  # 1 mb
                
                # we go through all the sheets
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    
                    # we go through all the lines
                    for row in ws.iter_rows(values_only=True):
                        for cell in row:
                            if cell is not None:
                                cell_str = str(cell)
                                text_parts.append(cell_str)
                                char_count += len(cell_str)
                                
                                # if you have typed enough text, check it
                                if char_count >= max_chars:
                                    combined_text = ' '.join(text_parts)
                                    match = keywords_pattern.search(combined_text)
                                    if match:
                                        print(f"keyword: {match}")
                                        wb.close()
                                        return True
                                    # reset to continue
                                    text_parts = []
                                    char_count = 0
                    
                    # checking the current sheet after processing
                    if text_parts:
                        combined_text = ' '.join(text_parts)
                        match = keywords_pattern.search(combined_text)
                        if match:
                            print(f"keyword: {match}")
                            wb.close()
                            return True
                        text_parts = []
                        char_count = 0
                
                wb.close()
                return False     
            except ImportError:
                print("openpyxl not installed. Run: pip install openpyxl")
                return False
            except Exception as e:
                print(f"Excel read error {file_path}: {e}")
                return False
        else:
            # unsupported format, skip
            return False
            
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return False

# peeks at if the file name contains at least one keyword.
def check_filename_for_keywords(filename: str, pattern: re.Pattern = KEYWORDS_PATTERN):
    # removing the file extension for verification
    name_without_ext = os.path.splitext(filename)[0]
    return bool(pattern.search(name_without_ext))

# finds files containing keywords.
def find_files_with_keywords(
    path: str,
    file_list: List[str], 
    keywords_pattern: re.Pattern = KEYWORDS_PATTERN
) -> List[str]:
    
    # step 1: look for files with keywords
    matched_files: List[str] = []
    
    for file_name in file_list:
        print(f"Checking: {os.path.basename(file_name)}")
        if (check_filename_for_keywords(file_name)):
            matched_files.append(file_name)
            print(f"Keywords found" )
            continue

        print(f"Checking file internals: {os.path.basename(file_name)}")
        if (check_keywords_in_file(path + '/' + file_name)):
            matched_files.append(file_name)
            print(f"Keywords found")


    print(matched_files)
    
    return matched_files

from database.db import get_all_from_processing_queue, update_processing_queue_field

# chews through all pdf files from the tender folder and stashes them in the database.if there are several pdfs in the folder, all their positions will be included in one tender (via upsert in save_tender_to_db).when restarted, old tender positions are deleted and written over again (see db.save_tender_to_db).
def import_pdf_files_from_folder_to_database(folder_path: str, tender_number: str = None, customer_name: str = None):
    if not os.path.exists(folder_path):
        print(f"Folder {folder_path} not found")
        return

    if tender_number is None:
        tender_number = os.path.basename(os.path.normpath(folder_path))

    files_in_folder = [
        f for f in os.listdir(folder_path)
        if os.path.isfile(os.path.join(folder_path, f))
    ]

    # we accumulate items from all pdfs of this tender in order to record them with one upsert
    all_items: List[dict] = []
    processed_pdfs: List[str] = []

    for file in files_in_folder:
        file_extension = os.path.splitext(file)[1].lower()
        if file_extension != '.pdf':
            continue

        pdf_path = os.path.join(folder_path, file)
        print(f"Parsing pdf: {pdf_path}")

        json_path = os.path.splitext(pdf_path)[0] + ".json"

        try:
            if os.path.exists(json_path):
                import json
                print(f"JSON already exists, loading: {json_path}")
                with open(json_path, 'r', encoding='utf-8') as jf:
                    parsed = json.load(jf)
            else:
                parsed = parse_pdf_to_json(pdf_path)

            items = parsed.get("items", []) or []
            # we mark which file each position came from
            for item in items:
                item.setdefault("_source_pdf", file)
            all_items.extend(items)
            processed_pdfs.append(file)
        except Exception as e:
            print(f"Failed to parse or load {pdf_path}: {e}")

    if not all_items:
        print(f"In folder {folder_path} no parsed positions, writing nothing to DB")
        return

    merged = {"items": all_items}
    pdf_source = ", ".join(processed_pdfs) if processed_pdfs else None

    db_id = save_tender_to_db(
        tender_number=tender_number,
        parsed_json=merged,
        pdf_source_file=pdf_source,
        customer_name=customer_name
    )

    if db_id is not None:
        print(f"Tender {tender_number} → r_luxai.tenders.id={db_id}, positions: {len(all_items)}")
    else:
        print(f"Saving tender {tender_number} to DB failed")

def get_file_hash(file_path: str, algorithm: str = "sha256") -> str:
    """
        Calculate file hash
    """
    hash_func = hashlib.new(algorithm)
    with open(file_path, "rb") as f:
        for chunk in iter(lambda: f.read(8192), b""):
            hash_func.update(chunk)
    return hash_func.hexdigest()

def get_all_hashes_in_folder(folder_path: str, algorithm: str = "sha256") -> set:
    """
        Return a set of file hashs in folder
    """
    hashes = set()
    for root, dirs, files in os.walk(folder_path):
        for file in files:
            file_path = os.path.join(root, file)
            try:
                file_hash = get_file_hash(file_path, algorithm)
                hashes.add(file_hash)
            except Exception as e:
                print(f"Error hashing {file_path}: {e}")
    return hashes

def move_all_files(source_folder: str, destination_folder: str, algorithm: str = "sha256"):
    print("\n\n\n")
    print(f"Move files from {source_folder} to {destination_folder}")


    if not os.path.exists(source_folder):
        print(f"Folders {source_folder} isn`t exists")
        return
    
    if not os.path.isdir(source_folder):
        print(f"{source_folder} isn`t folders")
        return
    
    if not os.path.exists(destination_folder):
        os.makedirs(destination_folder)

    # get hashes in dest_folder
    print(f"Scanning destination folder for existing hashes...")
    existing_hashes = get_all_hashes_in_folder(destination_folder, algorithm)
    print(f"Found {len(existing_hashes)} unique file hashes in destination")

    # statistic
    moved_count = 0
    skipped_count = 0
    deleted_count = 0
    error_count = 0

    # processing every file in source_folder
    for file in os.listdir(source_folder):
        source_path = os.path.join(source_folder, file)
        
        # skip if isn`t file
        if not os.path.isfile(source_path):
            continue
        
        try:
            # calculate file hash
            file_hash = get_file_hash(source_path, algorithm)
            
            # check hash in dest_folder
            if file_hash in existing_hashes:
                # duplicate - skip file
                print(f"File {source_path} has copy in {destination_folder} -> skip this file")
                os.remove(source_path)
                deleted_count += 1
                print(f"Deleted duplicate: {file}")
            else:
                # new file - move
                dest_path = os.path.join(destination_folder, file)
                
                # if file has that name - rename
                if os.path.exists(dest_path):
                    base, ext = os.path.splitext(file)
                    counter = 1
                    while os.path.exists(dest_path):
                        new_name = f"{base}_{counter}{ext}"
                        dest_path = os.path.join(destination_folder, new_name)
                        counter += 1
                    print(f"Renamed: {file} -> {os.path.basename(dest_path)}")
                
                # move file
                shutil.move(source_path, dest_path)
                moved_count += 1
                
                # add hash to set
                existing_hashes.add(file_hash)
                print(f"Moved: {file}")
                
        except Exception as e:
            print(f"Error processing {file}: {e}")
            error_count += 1

        # Удаляем исходную папку (должна быть пустой)
    try:
        if os.path.exists(source_folder):
            os.rmdir(source_folder)  # удаляем только пустую папку
            print(f"Removed source folder: {source_folder}")
    except OSError:
        print(f"Source folder not empty or could not be removed: {source_folder}")
    except Exception as e:
        print(f"Could not remove source folder {source_folder}: {e}")

def filter_single_tender_files(tender: dict):
    """
    Filters and converts raw files in a single tender's folder down to just the relevant PDFs.

    @param tender Dictionary containing tender data from the processing queue.
    """
    if not tender['files_downloaded'] or tender['files_filtered']:
        return

    full_tender_num = tender['tender_number']
    number = full_tender_num[2:] if len(full_tender_num) > 2 else full_tender_num

    print(f"Examining files for tender {number}")

    folder_path: str = os.path.join("./tenders_files", number)
    if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
        return

    files_in_folder = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]

    files_with_keywords = find_files_with_keywords(folder_path, files_in_folder)

    if (files_with_keywords):   # if we find files with the necessary words, then we convert them and delete all the others
        convert_files_to_pdf(folder_path, files_with_keywords)
        delete_files(folder_path, files_with_keywords)

    rename_all_files_in_folder(folder_path, number)

    source_folder = folder_path
    destination_folder = "backend/webui/static/webui/files"
    move_all_files(source_folder, destination_folder)

    tender['files_filtered'] = True
    update_processing_queue_field(full_tender_num, 'files_filtered', True)
    print(f"Marked {full_tender_num} as filtered")

def parse_single_tender_files(tender: dict):
    """
    Triggers the parsing of a single tender's processed PDFs.

    @param tender Dictionary containing tender data from the processing queue.
    """
    if not tender.get('files_filtered') or tender.get('files_parsed'):
        return

    full_tender_num = tender['tender_number']
    number = full_tender_num[2:] if len(full_tender_num) > 2 else full_tender_num
    customer_name = tender['customer']

    folder_path: str = os.path.join("../backend/webui/static/webui/files", number)
    if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
        return

    # pass the full tender number and customer name to the database saving function
    import_pdf_files_from_folder_to_database(folder_path, tender_number=str(full_tender_num), customer_name=customer_name)

    tender['files_parsed'] = True
    update_processing_queue_field(full_tender_num, 'files_parsed', True)
    print(f"Marked {full_tender_num} as parsed and completed")

def file_filter():
    """
    Iterates through the queue and processes raw files for filtering and conversion to PDF.
    """
    tenders = get_all_from_processing_queue()

    for tender in tenders:
        filter_single_tender_files(tender)

def file_parser():
    """
    Iterates through the queue and parses the finalized PDF documents into JSON.
    """
    tenders = get_all_from_processing_queue()

    for tender in tenders:
        parse_single_tender_files(tender)

if __name__ == "__main__":
    from database.db import deduplicate_tenders_in_db

    deduplicate_tenders_in_db()
    file_filter()
