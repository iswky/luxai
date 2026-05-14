from typing import List, Tuple, Any




from llm import parse_pdf_to_json
from db import save_tender_to_db
import os
import re


# grabs keywords from a text file.file format: each keyword on a new line.empty lines and lines starting with # are ignored.
def load_keywords_from_file(filename: str = "file_filter_keywords.txt") -> List[str]:
    keywords = []
    
    # checking the existence of the file
    if not os.path.exists(filename):
        print(f"File {filename} not found. Using default keywords.")
        return get_default_keywords()
    
    try:
        with open(filename, 'r', encoding='utf-8') as f:
            for line_num, line in enumerate(f, 1):
                line = line.strip()
                
                # skip empty lines and comments
                if not line or line.startswith('#'):
                    continue
                
                keywords.append(line)
        
        if not keywords:
            print(f"File {filename} contains no keywords. Using defaults.")
            return get_default_keywords()
        
        print(f"Loaded {len(keywords)} keywords from {filename}")
        return keywords
        
    except Exception as e:
        print(f"Error reading {filename}: {e}")
        print("Using default keywords.")
        return get_default_keywords()

# spits out a standard list of keywords (if the file is not found).
def get_default_keywords() -> List[str]:
    return [
        "описание объекта закупки",
        "ООЗ",
        "наименование объекта закупки",
        "технические характеристики",
        "ТЗ",
        "техническое задание",
        "требования к товару",
        "характеристики товара",
        "функциональные характеристики",
        "показатели товара",
        "спецификация",
        "перечень поставляемого товара",
        "соответствие товара",
        "КТРУ",
        "каталог товаров работ услуг"
    ]

keywords: List[str] = load_keywords_from_file()

# regular expression for quick search (case insensitive)
KEYWORDS_PATTERN = re.compile(
    r"(" + "|".join(re.escape(kw) for kw in keywords) + r")",
    re.IGNORECASE
)

# a simple version of deleting files.
def delete_files(path: str, filenames: List[str]):
    for filename in filenames:
        file_path = os.path.join(path, filename)

        file_extension = os.path.splitext(filename)[1].lower()

        if file_extension in ['.pdf']:
            continue

        try:
            if os.path.isfile(file_path):
                os.remove(file_path)
                print(f"Deleted: {filename}")
            else:
                print(f"File not found: {filename}")
        except Exception as e:
            print(f"Error deleting {filename}: {e}")

def convert_files_to_pdf(folder_path: str, files: List[str]):
    # checking the existence of the directory
    if not os.path.exists(folder_path):
        raise FileNotFoundError(f"Директория '{folder_path}' не существует")
    
    if not os.path.isdir(folder_path):
        raise NotADirectoryError(f"'{folder_path}' не является директорией")

    for filename in files:
        file_path = os.path.join(folder_path, filename)

        convert1file2pdf(file_path)

def convert1file2pdf(file_path: str):
    # checking the existence of the directory
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Директория '{file_path}' не существует")

    file_extension = os.path.splitext(file_path)[1].lower()
    output_filename = os.path.splitext(file_path)[0] + ".pdf"
    
    try:               
        if file_extension in ['.pdf']:
            return
        elif file_extension in ['.xlsx', '.xls']:
            convert_excel_to_pdf(file_path, output_filename)
        elif file_extension in ['.docx', '.doc']:
            convert_word_to_pdf(file_path, output_filename)
        else:
            print(f"Unsupported format: {file_path}")
            return
        
        print(f"Converted: {file_path} -> {output_filename}")
        
    except Exception as e:
        print(f"Conversion error {file_path}: {str(e)}")
    


# converts the specified files to pdf format.supports: .txt, .docx, .doc, .jpg, .jpeg, .png, .html
def convert_to_pdf(path: str, files: List[str]):
    # checking the existence of the directory
    if not os.path.exists(path):
        raise FileNotFoundError(f"Директория '{path}' не существует")
    
    if not os.path.isdir(path):
        raise NotADirectoryError(f"'{path}' не является директорией")
    
    converted_files = []
    failed_files = []
    
    for filename in files:
        file_path = os.path.join(path, filename)
        # print(file_path)
        
        # checking the existence of the file
        if not os.path.exists(file_path):
            print(f"File not found: {filename}")
            failed_files.append(filename)
            continue
        
        # getting the file extension
        file_extension = os.path.splitext(filename)[1].lower()
        output_filename = os.path.splitext(filename)[0] + ".pdf"
        output_path = os.path.join(path, output_filename)
        
        try:               
            if file_extension in ['.pdf']:
                continue
            if file_extension in ['.xlsx', '.xls']:
                convert_excel_to_pdf(file_path, output_path)
            if file_extension in ['.docx']:
                convert_word_to_pdf(file_path, output_path)
            else:
                print(f"Unsupported format: {filename}")
                failed_files.append(filename)
                continue
            
            print(f"Converted: {filename} -> {output_filename}")
            converted_files.append(filename)
            
        except Exception as e:
            print(f"Conversion error {filename}: {str(e)}")
            failed_files.append(filename)
    
    # statistics output
    print("\n" + "="*50)
    print(f"Conversion stats:")
    print(f"Success: {len(converted_files)} files")
    print(f"Errors: {len(failed_files)} files")
    
    if failed_files:
        print(f"\n Failed to convert: {', '.join(failed_files)}")

# helper funcs for different formats:
# converts word document to pdf
def convert_word_to_pdf(input_path: str, output_path: str = None):

    import subprocess

    # if output_path is not specified, create next to the source file
    if output_path is None:
        output_path = os.path.splitext(input_path)[0] + '.pdf'
    
    # we get the directory for output
    output_dir = os.path.dirname(output_path)
    if not output_dir:
        output_dir = '.'
    
    # conversion via libreoffice
    cmd = [
        'libreoffice',
        '--headless',           # no gui
        '--convert-to', 'pdf',  # convert to pdf
        '--outdir', output_dir, # output directory
        input_path              # input file
    ]
    
    try:
        result = subprocess.run(cmd, capture_output=True, text=True, check=True)
        
        # libreoffice stashes a file with the same name in output_dir
        generated_pdf = os.path.join(output_dir, os.path.basename(input_path).replace('.docx', '.pdf'))
        
        # if you need to rename or move
        if generated_pdf != output_path:
            os.rename(generated_pdf, output_path)
        
        print(f"Conversion successful: {output_path}")
        return output_path
        
    except subprocess.CalledProcessError as e:
        print(f"Conversion error: {e}")
        print(f"Stderr: {e.stderr}")
        raise
    except FileNotFoundError:
        raise Exception("LibreOffice не установлен. Установите: sudo apt install libreoffice")

# converts excel to pdf with correct data reading
def convert_excel_to_pdf(input_path: str, output_path: str):
    import pandas as pd
    import matplotlib.pyplot as plt
    from matplotlib.backends.backend_pdf import PdfPages

    # print(input_path)
    
    # read with additional params
    df = pd.read_excel(
        input_path,
        header=None,  # don't use the first line as a title
        keep_default_na=False,  # don't replace empty values ​​with nan
        na_filter=False  # disable na filtering
    )
    
    # removing completely empty rows and columns
    df = df.dropna(how='all', axis=0)  # removing empty lines
    df = df.dropna(how='all', axis=1)  # removing empty columns
    
    print(f"Table size: {df.shape}")
    
    if df.empty or df.shape[0] == 0 or df.shape[1] == 0:
        raise ValueError("Excel файл не содержит данных")
    
    # replace nan with empty strings
    df = df.fillna('')
    
    # calculate the size
    num_rows, num_cols = df.shape
    fig_width = max(12, num_cols * 2)
    fig_height = max(8, num_rows * 0.6)
    
    fig, ax = plt.subplots(figsize=(fig_width, fig_height))
    ax.axis('tight')
    ax.axis('off')
    
    # create a table without headers (since header=none)
    table = ax.table(
        cellText=df.values,
        cellLoc='left',  # left alignment
        loc='center',
        colWidths=[1.0/num_cols] * num_cols
    )
    
    table.auto_set_font_size(False)
    table.set_fontsize(9)
    table.scale(1, 2)
    
    with PdfPages(output_path) as pdf:
        pdf.savefig(fig, bbox_inches='tight', dpi=150)
    
    plt.close()
    print(f"PDF saved: {output_path}")

# nukes all files in the specified directory except those specified in neccesary_files.
def delete_extra_files(path: str, Neccesary_files: List[str]):
    # checking if the directory exists
    if not os.path.exists(path):
        raise FileNotFoundError(f"Директория '{path}' не существует")
    
    # checking if the path is a directory
    if not os.path.isdir(path):
        raise NotADirectoryError(f"'{path}' не является директорией")
    
    # get a list of all files in a directory
    all_files = [f for f in os.listdir(path) if os.path.isfile(os.path.join(path, f))]
    
    print(all_files)

    # we create many necessary files for quick search
    necessary_set = set(Neccesary_files)
    
    # removing unnecessary files
    deleted_count = 0
    for file in all_files:
        if file not in necessary_set:
            print(f"deleting {file}")
            file_path = os.path.join(path, file)
            try:
                os.remove(file_path)
                print(f"Deleted: {file}")
                deleted_count += 1
            except Exception as e:
                print(f"Error deleting '{file}': {e}")
    
    print(f"\nDeleted files: {deleted_count}")
    print(f"Saved files: {len(necessary_set & set(all_files))}")

# peeks at if the file contains at least one keyword.supported formats: .txt, .docx, .pdf, .xlsx, .odt, .ods, .rtf
def check_keywords_in_file(file_path: str, keywords_pattern: re.Pattern = KEYWORDS_PATTERN) -> bool:
    try:
        ext = os.path.splitext(file_path)[1].lower()
        
        # text files and html
        if ext in ['.txt', '.html', '.htm', '.xml', '.csv', '.md', '.rst']:
            with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                content = f.read(1024 * 1024)  # read the first 1mb for speed
            return bool(keywords_pattern.search(content))
        
        # docx
        elif ext == '.docx':
            try:
                from docx import Document
                doc = Document(file_path)
                text = ' '.join([para.text for para in doc.paragraphs])
                # also check the tables
                for table in doc.tables:
                    for row in table.rows:
                        for cell in row.cells:
                            text += ' ' + cell.text
                match = keywords_pattern.search(text[:1024*1024])
                if match:
                    print(f"keyword: {match.group(0)}")
                    return True
                else:
                    return False
            except ImportError:
                print("python-docx not installed. Run: pip install python-docx")
                return False
        
        # pdf
        elif ext == '.pdf':
            try:
                from pypdf import PdfReader
                reader = PdfReader(file_path)
                text = ''
                for page in reader.pages[:10]:  # first 10 pages
                    text += page.extract_text() or ''
                return bool(keywords_pattern.search(text[:1024*1024]))
            except ImportError:
                print("pypdf not installed. Run: pip install pypdf")
                return False
        
        # excel
        elif ext in ['.xlsx', '.xls']:
            try:

                
                # loading the workbook (.xlsx only, .xls is not supported by openpyxl)
                wb = load_workbook(file_path, read_only=True, data_only=True)
                
                text_parts = []
                char_count = 0
                max_chars = 1024 * 1024  # 1 mb
                
                # we go through all the sheets
                for sheet_name in wb.sheetnames:
                    ws = wb[sheet_name]
                    
                    # we go through all the lines
                    for row in ws.iter_rows(values_only=True):
                        for cell in row:
                            if cell is not None:
                                cell_str = str(cell)
                                text_parts.append(cell_str)
                                char_count += len(cell_str)
                                
                                # if you have typed enough text, check it
                                if char_count >= max_chars:
                                    combined_text = ' '.join(text_parts)
                                    match = keywords_pattern.search(combined_text)
                                    if match:
                                        print(f"keyword: {match}")
                                        wb.close()
                                        return True
                                    # reset to continue
                                    text_parts = []
                                    char_count = 0
                    
                    # checking the current sheet after processing
                    if text_parts:
                        combined_text = ' '.join(text_parts)
                        match = keywords_pattern.search(combined_text)
                        if match:
                            print(f"keyword: {match}")
                            wb.close()
                            return True
                        text_parts = []
                        char_count = 0
                
                wb.close()
                return False     
            except ImportError:
                print("openpyxl not installed. Run: pip install openpyxl")
                return False
            except Exception as e:
                print(f"Excel read error {file_path}: {e}")
                return False
        else:
            # unsupported format, skip
            return False
            
    except Exception as e:
        print(f"Error reading {file_path}: {e}")
        return False

# peeks at if the file name contains at least one keyword.
def check_filename_for_keywords(filename: str, pattern: re.Pattern = KEYWORDS_PATTERN):
    # removing the file extension for verification
    name_without_ext = os.path.splitext(filename)[0]
    return bool(pattern.search(name_without_ext))

# finds files containing keywords.
def find_files_with_keywords(
    path: str,
    file_list: List[str], 
    keywords_pattern: re.Pattern = KEYWORDS_PATTERN
) -> List[str]:
    
    # step 1: look for files with keywords
    matched_files: List[str] = []
    
    for file_name in file_list:
        print(f"Checking: {os.path.basename(file_name)}")
        if (check_filename_for_keywords(file_name)):
            matched_files.append(file_name)
            print(f"Keywords found" )
            continue

        print(f"Checking file internals: {os.path.basename(file_name)}")
        if (check_keywords_in_file(path + '/' + file_name)):
            matched_files.append(file_name)
            print(f"Keywords found")


    print(matched_files)
    
    return matched_files

from db import get_all_from_processing_queue, update_processing_queue_field

# chews through all pdf files from the tender folder and stashes them in the database.if there are several pdfs in the folder, all their positions will be included in one tender (via upsert in save_tender_to_db).when restarted, old tender positions are deleted and written over again (see db.save_tender_to_db).
def import_pdf_files_from_folder_to_database(folder_path: str, tender_number: str = None, customer_name: str = None):
    if not os.path.exists(folder_path):
        print(f"Folder {folder_path} not found")
        return

    if tender_number is None:
        tender_number = os.path.basename(os.path.normpath(folder_path))

    files_in_folder = [
        f for f in os.listdir(folder_path)
        if os.path.isfile(os.path.join(folder_path, f))
    ]

    # we accumulate items from all pdfs of this tender in order to record them with one upsert
    all_items: List[dict] = []
    processed_pdfs: List[str] = []

    for file in files_in_folder:
        file_extension = os.path.splitext(file)[1].lower()
        if file_extension != '.pdf':
            continue

        pdf_path = os.path.join(folder_path, file)
        print(f"Parsing pdf: {pdf_path}")

        try:
            parsed = parse_pdf_to_json(pdf_path)
            items = parsed.get("items", []) or []
            # we mark which file each position came from
            for item in items:
                item.setdefault("_source_pdf", file)
            all_items.extend(items)
            processed_pdfs.append(file)
        except Exception as e:
            print(f"Failed to parse {pdf_path}: {e}")

    if not all_items:
        print(f"In folder {folder_path} no parsed positions, writing nothing to DB")
        return

    merged = {"items": all_items}
    pdf_source = ", ".join(processed_pdfs) if processed_pdfs else None

    db_id = save_tender_to_db(
        tender_number=tender_number,
        parsed_json=merged,
        pdf_source_file=pdf_source,
        customer_name=customer_name
    )

    if db_id is not None:
        print(f"Tender {tender_number} → r_luxai.tenders.id={db_id}, positions: {len(all_items)}")
    else:
        print(f"Saving tender {tender_number} to DB failed")

def rename_all_files_in_folder(folder_path: str, tender_id: str):
    if not os.path.exists(folder_path):
        print(f"Folder {folder_path} not found")
        return

    files_in_folder = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]

    i = 1
    for file in files_in_folder:
        file_ext = os.path.splitext(file)[1]
        old_path = os.path.join(folder_path, file)
        new_path = os.path.join(folder_path, f"{tender_id}_{i}{file_ext}")

        os.rename(old_path, new_path)
        i += 1

def file_filter():
    tenders = get_all_from_processing_queue()

    for tender in tenders:
        if not tender['files_downloaded'] or tender['files_filtered']:
            continue

        full_tender_num = tender['tender_number']
        number = full_tender_num[2:] if len(full_tender_num) > 2 else full_tender_num
        customer_name = tender['customer']

        print(f"Examining files for tender {number}")

        folder_path: str = os.path.join("./tenders_files", number)
        if not os.path.exists(folder_path) or not os.path.isdir(folder_path):
            continue

        files_in_folder = [f for f in os.listdir(folder_path) if os.path.isfile(os.path.join(folder_path, f))]

        files_with_keywords = find_files_with_keywords(folder_path, files_in_folder)


        if (not files_with_keywords):   # if we haven’t found a single file with the words we need, then we convert all the files
            convert_files_to_pdf(folder_path, files_in_folder)
        else:   # if we find files with the necessary words, then we convert them and delete all the others
            convert_files_to_pdf(folder_path, files_with_keywords)

        delete_files(folder_path, files_in_folder)

        rename_all_files_in_folder(folder_path, number)

        # pass the full tender number and customer name to the database saving function
        import_pdf_files_from_folder_to_database(folder_path, tender_number=str(full_tender_num), customer_name=customer_name)

        update_processing_queue_field(full_tender_num, 'files_filtered', True)
        print(f"Marked {full_tender_num} as filtered and completed")


if __name__ == "__main__":
    from db import deduplicate_tenders_in_db

    deduplicate_tenders_in_db()
    file_filter()
